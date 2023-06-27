import lxml.html as html
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from ebooklib import epub, utils
from googletrans import Translator


def translation_func(text):
    translator = Translator()
    result = translator.translate(text, dest='ru')
    return result.text


def open_epub():
    class_exeption = ['hljs-built_in', 'packt_figref', 'hljs-con-built_in', 'hljs-keyword-slc', 'hljs-number-slc',
                      'code-highlight', 'hljs-tag', 'hljs-built_in-slc', 'hljs-name', 'hljs-meta', 'hljs-code',
                      'hljs-slc', 'hljs-number', 'hljs-string-slc', 'python', 'keyWord', 'keystroke', 'italic',
                      'chapterNumber', 'bulletList', 'hljs-title', 'hljs-params', 'normal', 'hljs-comment',
                      'hljs-comment-slc', 'inlineCode', 'hljs-title-slc', 'hljs-name-slc', 'chapterTitle', 'hljs-class',
                      'hljs-function', 'hljs-attr', 'url', 'heading-1', 'hljs-literal-slc', 'hljs-con-string',
                      'mediaobject', 'hljs-params-slc', 'heading-2', 'programlisting code', 'numberedList',
                      'hljs-literal', 'hljs-con', 'hljs-con-number', 'screenText', 'hljs-function-slc', 'hljs-con-meta',
                      'hljs-con-keyword', 'hljs-meta-keyword', 'hljs-string', 'hljs-keyword', 'codeHighlighted',
                      'Basic-Text-Frame', 'chapterRef', 'hljs-tag-slc', 'hljs-attr-slc', 'programlisting con', 'note']

    tag_exeption = ["code", 'a', 'strong', 'pre', 'span', 'html',
                    'div', 'body', "head"]

    book = epub.read_epub('book.epub')

    for item in book.get_items():

        if item.get_type() == 9:
            try:
                print('TYPE : ', item.get_type())
            except:
                print('TYPE : ', None)
            print('----------------------------------')
            try:
                print('NAME : ', item.get_name())
            except:
                print('NAME : ', None)
            print('----------------------------------')
            try:
                print('ID : ', item.get_id())
            except:
                print('ID : ', None)
            print('----------------------------------')

            soup = BeautifulSoup(item.get_content(), features="xml")
            count = 0
            total_string = len(list(soup.descendants))
            print(total_string)
            for child in soup.descendants:
                if child.name and child.string:
                    print(child.name, '->', child.attrs)

            #     if child.name == 'text':
            #
            #         tag_text_before = child.string
            #         translation_text = translation_func(tag_text_before)
            #         child.string = translation_text
            #
            #         count += 1
            #         print(tag_text_before, '->', translation_text, '\n', total_string, '->>', count)
            #
            #
            # item.set_content(soup.encode())
            # contents = utils.parse_string(item.get_content())
            # html.open_in_browser(contents)

        elif item.get_type() == -2:
            try:
                print('ITEM : ', item.get_content())
                style = item.get_content()
                nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css",
                                        content=style)
                book.add_item(nav_css)
            except:
                print('NAME : ', None)
        elif item.get_id() == "Chapter_66":

            try:
                print('NAME : ', item.get_name())
            except:
                print('NAME : ', None)
            print('----------------------------------')
            try:
                print('ID : ', item.get_id())
            except:
                print('ID : ', None)
            print('----------------------------------')
            try:
                print('ITEM : ', item.get_type())
            except:
                print('NAME : ', None)
            print(item.get_content())

            soup = BeautifulSoup(item.get_content(), features="xml")
            count = 0

            total_string = len(list(soup.descendants))
            # class_array = set()
            for child in soup.descendants:
                # print(child.name, '->', child.attrs)
                # if child.name and child.string:
                #     print(child.name, '->', child.attrs)
                #     #print(child.string)
                #     print(child.name, '->>', child.attrs.get('class'))
                #     print('----------------------------------')
                # if child.name == 'span':
                #     print(child.string, '->>', child.attrs.get('class'))
                # class_array.add(child.attrs.get('class'))
                # print()
                if not child.name in tag_exeption and child.name and child.string:  # and count < 10:

                    if child.parent.name == 'code':
                        continue

                    tag_text_before = child.string
                    translation_text = translation_func(tag_text_before)
                    child.string = translation_text
                    count += 1
                    print(total_string, '->>', count)
                #
                elif not child.name in tag_exeption and child.name:  # and count < 10:
                    new_contents = []
                    class_attr = child.attrs.get('class')
                    for content in child.contents:
                        if content.string and content.string not in ['\n', ' '] and not content.name:
                            tag_text_before = content.string
                            translation_text = translation_func(tag_text_before)
                            content = NavigableString(translation_text)
                        new_contents.append(content)
                        new_contents.append(" ")
                        count += 1
                    child.clear()
                    child.extend(new_contents)
                    child['class'] = class_attr
            #         print(total_string, '->>', count)
            item.set_content(soup.encode())
            #
            contents = utils.parse_string(item.get_content())
            html.open_in_browser(contents)

            print('==================================')
    # styles = book.get_items_of_type(ebooklib.ITEM_STYLE)
    # for style in styles:
    #     print(style)
    #     style = item.get_content()
    # nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css",
    #                              content=style)
    #
    # book.add_item(nav_css)
    # book.spine = ['nav', item]
    epub.write_epub('book.epub', book, {})


def open_document():
    document = Document("Django 4 By Example 2022[072-117].docx")
    style_set_func(document)
    # traslation_document(document)


def traslation_document(document):
    styles_exceptions = ['Para 03', 'Para 04', 'Para 17', 'Para 14', 'Para 22']
    runs_exceptions = ['00 Text', '09 Text', '15 Text', '05 Text', '07 Text', '10 Text', '11 Text', '22 Text',
                       '02 Text', '44 Text', '19 Text', '33 Text', '37 Text', '13 Text', '01 Text', '14 Text',
                       '36 Text', '28 Text', '26 Text', '29 Text', '24 Text', '45 Text', '25 Text', '16 Text',
                       '21 Text', '42 Text', '51 Text', '17 Text', '23 Text', '08 Text', '57 Text', '12 Text',
                       '30 Text']

    for paragraph in document.paragraphs:
        if not paragraph.style.name in styles_exceptions:

            # print(paragraph.style.name, " - ", translation_func(paragraph.text))

            for run in paragraph.runs:
                if not run.style.name in runs_exceptions:
                    run_text_befo = run.text
                    translation_text = translation_func(run.text) + ' '
                    run.clear()
                    run.add_text(translation_text)
                    # print("   ", run_text_befo, " - ", run.text)
        # styles.append(paragraph.style.name)
    document.save('Chapter_1_tran.docx')


def style_set_func(document):
    style_set = []
    runs_style_set = []
    for paragraph in document.paragraphs:
        # if not paragraph.style.name in style_set:
        print("#", paragraph.style.name, " - ", paragraph.text[:25])
        style_set.append(paragraph.style.name)

        for run in paragraph.runs:
            if run.style:
                # if not run.style.name in runs_style_set:
                print("****", run.style.name, " - ", run.text)
                runs_style_set.append(run.style.name)

    print(style_set)
    print(runs_style_set)


def main():
    open_epub()
    # open_document()


if __name__ == "__main__":
    # translator = Translator()
    # result = translator.translate("Adding comments to the post detail view", dest='ru')
    # print(result.text)
    main()
